"""
Document processing and file handling.
"""

import asyncio
import logging
from typing import List, Optional
from pathlib import Path
import aiofiles
import pypdf
from docx import Document as DocxDocument
from PIL import Image

from ..models import Document, DocumentMetadata, FileType, DocumentStatus, LanguageCode
from ..utils.settings import settings
from .ocr import ocr_processor
from .chunking import text_chunker

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document types and extract content."""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_directory)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def validate_file(self, file_path: Path) -> bool:
        """Validate file type and size."""
        try:
            # Check file exists
            if not file_path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # Check file size
            file_size = file_path.stat().st_size
            max_size = settings.max_file_size_mb * 1024 * 1024
            
            if file_size > max_size:
                logger.error(f"File too large: {file_size} bytes (max: {max_size})")
                return False
            
            # Check file extension
            file_extension = file_path.suffix.lower().lstrip('.')
            if file_extension not in settings.allowed_file_types_list:
                logger.error(f"Unsupported file type: {file_extension}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            return False
    
    def get_file_type(self, file_path: Path) -> FileType:
        """Determine file type from extension."""
        extension = file_path.suffix.lower().lstrip('.')
        
        type_mapping = {
            'pdf': FileType.PDF,
            'docx': FileType.DOCX,
            'txt': FileType.TXT,
            'jpg': FileType.JPG,
            'jpeg': FileType.JPEG,
            'png': FileType.PNG
        }
        
        return type_mapping.get(extension, FileType.TXT)
    
    async def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        # Try to extract text directly
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            text_content.append(page_text)
                        else:
                            # If no text extracted, it might be a scanned PDF
                            logger.info(f"No text found on page {page_num + 1}, may need OCR")
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""
    
    async def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""
    
    async def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
                return content
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                    return content
            except Exception as e:
                logger.error(f"TXT text extraction failed with encoding issues: {e}")
                return ""
        except Exception as e:
            logger.error(f"TXT text extraction failed: {e}")
            return ""
    
    async def extract_text_from_image(self, file_path: Path) -> tuple[str, float, LanguageCode]:
        """Extract text from image using OCR."""
        try:
            ocr_result = await ocr_processor.process_image(str(file_path))
            return ocr_result.text, ocr_result.confidence, ocr_result.language
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return "", 0.0, LanguageCode.ENGLISH
    
    async def process_document(self, file_path: Path) -> Document:
        """Process a document and extract its content."""
        # Validate file
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        # Create document metadata
        file_type = self.get_file_type(file_path)
        file_size = file_path.stat().st_size
        
        metadata = DocumentMetadata(
            filename=file_path.name,
            original_path=str(file_path),
            file_type=file_type,
            file_size=file_size,
            processing_status=DocumentStatus.PROCESSING
        )
        
        document = Document(metadata=metadata)
        
        try:
            # Extract content based on file type
            if file_type == FileType.PDF:
                content = await self.extract_text_from_pdf(file_path)
                
                # If PDF has little or no text, it might be scanned
                if len(content.strip()) < 100:
                    logger.info("PDF appears to be scanned, attempting OCR")
                    # TODO: Convert PDF pages to images and run OCR
                    # This would require additional libraries like pdf2image
                    
            elif file_type == FileType.DOCX:
                content = await self.extract_text_from_docx(file_path)
                
            elif file_type == FileType.TXT:
                content = await self.extract_text_from_txt(file_path)
                
            elif file_type in [FileType.JPG, FileType.JPEG, FileType.PNG]:
                content, confidence, language = await self.extract_text_from_image(file_path)
                metadata.ocr_confidence = confidence
                metadata.language = language
                
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Update document with extracted content
            document.content = content
            
            # Create chunks if content is available
            if content and content.strip():
                chunks = await text_chunker.create_chunks(document.id, content)
                document.chunks = chunks
                metadata.processing_status = DocumentStatus.COMPLETED
            else:
                logger.warning(f"No content extracted from {file_path}")
                metadata.processing_status = DocumentStatus.FAILED
                metadata.error_message = "No text content found"
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {e}")
            metadata.processing_status = DocumentStatus.FAILED
            metadata.error_message = str(e)
        
        return document

# Global document processor instance
document_processor = DocumentProcessor()